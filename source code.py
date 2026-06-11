# ----------------------------- Part 1: Imports and Initialization -----------------------------
import numpy as np
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from scipy.signal import savgol_filter
from docx import Document
from docx.shared import Inches
import os
import tempfile

# Universal Constants
q = 1.602e-19  # Electron charge (C)
k = 1.381e-23  # Boltzmann constant (J/K)
Aeff = 1e-4  # Effective diode area in cm²


class MainApplication:
    def __init__(self, root):
        self.root = root
        self.root.title("Semiconductor Analysis Toolkit")
        self.root.geometry("1350x800")

        # Create container frame
        self.container = tk.Frame(root)
        self.container.pack(side="top", fill="both", expand=True)
        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)

        # Initialize frames
        self.frames = {}
        for F in (IVAnalyzerApp, SurfaceStateAnalyzer):
            frame = F(self.container, self)
            self.frames[F] = frame
            frame.grid(row=0, column=0, sticky="nsew")

        self.show_frame(IVAnalyzerApp)

        # Add navigation buttons to both frames
        self.add_navigation_buttons()

    def show_frame(self, cont):
        frame = self.frames[cont]
        frame.tkraise()

    def add_navigation_buttons(self):
        # Add buttons to both frames with specific positioning
        iv_frame = self.frames[IVAnalyzerApp]
        nav_btn_iv = tk.Button(iv_frame, text="Go to Surface State Analyzer",
                               command=lambda: self.show_frame(SurfaceStateAnalyzer))
        nav_btn_iv.place(relx=0.01, rely=0.98, anchor="sw")  # Bottom left corner

        ss_frame = self.frames[SurfaceStateAnalyzer]
        nav_btn_ss = tk.Button(ss_frame, text="Go to I-V Analyzer",
                               command=lambda: self.show_frame(IVAnalyzerApp))
        nav_btn_ss.place(relx=0.01, rely=0.98, anchor="sw")  # Bottom center (original position)

class IVAnalyzerApp(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.controller = controller
        self.df = None
        self.filename = None
        self.smoothing_method = None
        self.method = tk.StringVar(value="Thermionic")
        self.diode_type = tk.StringVar(value="n-type")  # Default is n-type

        self.setup_gui()

    def setup_gui(self):
        # ---------------- Left Panel ----------------
        frame_left = tk.Frame(self)
        frame_left.pack(side=tk.LEFT, padx=10, pady=10, anchor="n")

        # Load File Button
        tk.Button(frame_left, text="📁 Load File", command=self.load_file).pack(pady=2)

        # Column selection
        col_frame = tk.Frame(frame_left)
        col_frame.pack(pady=5)
        tk.Label(col_frame, text="Voltage Col:").grid(row=0, column=0)
        self.volt_entry = tk.Entry(col_frame, width=5)
        self.volt_entry.grid(row=0, column=1)
        tk.Label(col_frame, text="Current Col:").grid(row=0, column=2)
        self.curr_entry = tk.Entry(col_frame, width=5)
        self.curr_entry.grid(row=0, column=3)

        # Temperature Entry
        tk.Label(frame_left, text="Temperature (K):").pack()
        self.temp_entry = tk.Entry(frame_left, width=10)
        self.temp_entry.insert(0, "300")  # Default 300K
        self.temp_entry.pack(pady=2)

        # Range
        tk.Label(frame_left, text="Vmin (V):").pack()
        self.vmin_entry = tk.Entry(frame_left, width=10)
        self.vmin_entry.pack(pady=2)
        tk.Label(frame_left, text="Vmax (V):").pack()
        self.vmax_entry = tk.Entry(frame_left, width=10)
        self.vmax_entry.pack(pady=2)

        # Smoothing
        tk.Button(frame_left, text="🧹 Choose Smoothing", command=self.choose_smoothing).pack(pady=4)

        # Method
        tk.Label(frame_left, text="Select Method:").pack()
        method_menu = ttk.Combobox(frame_left, values=["Thermionic", "Cheung"], textvariable=self.method,
                                   state="readonly")
        method_menu.pack()
        method_menu.bind("<<ComboboxSelected>>", self.on_method_change)

        # Diode Type
        tk.Label(frame_left, text="Diode Type (Richardson constant):").pack()
        type_menu = ttk.Combobox(frame_left, values=["n-type", "p-type"], textvariable=self.diode_type,
                                 state="readonly")
        type_menu.pack()

        # Analysis Buttons
        tk.Button(frame_left, text="📊 Analyze", command=self.analyze).pack(pady=5)
        tk.Button(frame_left, text="🧠 Best Exponential Region", command=self.best_exponential_region).pack(pady=5)
        tk.Button(frame_left, text="📝 Save to Word", command=self.save_to_word).pack(pady=5)

        # ---------------- Plot Canvas ----------------
        self.canvas_frame = tk.Frame(self)
        self.canvas_frame.pack(side=tk.TOP, pady=10)

        # ---------------- Results & Equations Panel ----------------
        self.result_frame = tk.Frame(self)
        self.result_frame.pack(side=tk.BOTTOM, fill=tk.BOTH, expand=True)

        self.result_text = tk.Text(self.result_frame, width=80, font=("Courier New", 10))
        self.result_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.equation_text = tk.Text(self.result_frame, width=60, font=("Courier New", 10))
        self.equation_text.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Configure text styles
        self.result_text.tag_config("header", font=("Courier New", 12, "bold"))
        self.result_text.tag_config("bold", font=("Courier New", 12, "bold"))
        self.equation_text.tag_config("header", font=("Courier New", 11, "bold"))

    def on_method_change(self, event=None):
        """Don't reset Vmin/Vmax when switching methods"""
        pass

    def load_file(self):
        """Load Excel or CSV file into DataFrame."""
        file_path = filedialog.askopenfilename(filetypes=[("Excel or CSV files", "*.xlsx *.xls *.csv")])
        if file_path:
            self.filename = file_path
            try:
                if file_path.endswith(".csv"):
                    self.df = pd.read_csv(file_path)
                else:
                    self.df = pd.read_excel(file_path)

                # Convert to numeric and drop non-numeric rows
                self.df = self.df.apply(pd.to_numeric, errors='coerce').dropna()
                messagebox.showinfo("Loaded", f"File loaded: {os.path.basename(file_path)}\nRows: {len(self.df)}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load file:\n{e}")
                self.df = None

    def clear_canvas(self):
        """Remove all plots from canvas frame."""
        for widget in self.canvas_frame.winfo_children():
            widget.destroy()

    def plot_on_canvas(self, fig):
        """Embed a Matplotlib figure in the canvas frame."""
        canvas = FigureCanvasTkAgg(fig, master=self.canvas_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(side=tk.LEFT, padx=10)

    def display_results(self, results, raw_data, method):
        """Display results in bold, large text with clean formatting."""
        self.result_text.delete("1.0", tk.END)
        self.equation_text.delete("1.0", tk.END)

        # Header
        self.result_text.insert(tk.END, f"{method} Method Results\n", "header")
        self.result_text.insert(tk.END, "-" * 50 + "\n", "header")

        for label, value in results:
            # Show normal decimal unless exponential is needed
            if abs(value) > 9999 or abs(value) < 1e-3:
                value_str = f"{value:.5e}"
            else:
                value_str = f"{value:.6f}"
            self.result_text.insert(tk.END, f"{label:<40}: {value_str}\n", "bold")

        self.result_text.insert(tk.END, "\nRaw Data (Used for Fit)\n", "header")
        self.result_text.insert(tk.END, "Voltage (V)       Current (A)\n", "header")
        self.result_text.insert(tk.END, "-" * 50 + "\n", "header")

        for v, i in raw_data:
            val_str = f"{i:.6e}" if abs(i) < 1e-3 else f"{i:.6f}"
            self.result_text.insert(tk.END, f"{v:<18.4f} {val_str:<}\n", "bold")

        # Equations
        self.equation_text.insert(tk.END, "Equations Used:\n\n", "header")
        if method == "Thermionic":
            self.equation_text.insert(tk.END,
                                      "n     = q / (slope × k × T)\n"
                                      "I₀    = exp(intercept)\n"
                                      "Φb    = (kT/q) × ln((A* × Aeff × T²) / I₀)\n"
                                      "ln(I) = slope × V + intercept\n")
        else:
            self.equation_text.insert(tk.END,
                                      "dV/dlnI = Rs1 × I + n(kT/q)\n"
                                      "n       = (slope1 × q) / (k × T)\n"
                                      "I₀      = exp(-q * phi_b / (n * k * T))\n"
                                      "H(I)    = V - (kT/q) × ln(I) = Rs2 × I + Φb\n"
                                      "Φb      = intercept2)\n")

    def save_to_word(self):
        """Export results and plots to a Word document."""
        if not hasattr(self, 'result_text') or not hasattr(self, 'canvas_frame'):
            messagebox.showerror("Error", "No analysis results to save")
            return

        try:
            doc = Document()
            doc.add_heading("I-V Analysis Report", level=1)

            # Results Section
            doc.add_heading("Results", level=2)
            for line in self.result_text.get("1.0", tk.END).splitlines():
                if line.strip():  # Skip empty lines
                    doc.add_paragraph(line)

            # Equations Section
            doc.add_page_break()
            doc.add_heading("Equations Used", level=2)
            for line in self.equation_text.get("1.0", tk.END).splitlines():
                if line.strip():  # Skip empty lines
                    doc.add_paragraph(line)

            # Add plots
            temp_dir = tempfile.gettempdir()
            for i, widget in enumerate(self.canvas_frame.winfo_children()):
                if isinstance(widget, FigureCanvasTkAgg):
                    fig = widget.figure
                    img_path = os.path.join(temp_dir, f"iv_plot_{i}.png")
                    fig.savefig(img_path, bbox_inches='tight', dpi=300)
                    doc.add_page_break()
                    doc.add_heading(f"Plot {i + 1}", level=3)
                    doc.add_picture(img_path, width=Inches(5.5))
                    os.remove(img_path)  # Clean up temp file

            # Save dialog
            path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                title="Save Report As"
            )
            if path:
                doc.save(path)
                messagebox.showinfo("Saved", f"Report saved to:\n{path}")
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save document:\n{e}")

    def choose_smoothing(self):
        """Popup to select smoothing method."""
        win = tk.Toplevel(self)
        win.title("Choose Smoothing Method")

        def set_savgol():
            self.smoothing_method = "savgol"
            win.destroy()

        def set_moving_avg():
            self.smoothing_method = "moving"
            win.destroy()

        tk.Button(win, text="Savitzky-Golay", command=set_savgol).pack(pady=5)
        tk.Button(win, text="Moving Average", command=set_moving_avg).pack(pady=5)

    def smooth(self, data):
        """Apply the selected smoothing algorithm."""
        if self.smoothing_method == "savgol":
            return savgol_filter(data, 11, 3)
        elif self.smoothing_method == "moving":
            return np.convolve(data, np.ones(10) / 10, mode='same')
        return data

    def get_A_star(self):
        """Return the Richardson constant based on diode type."""
        return 120 if self.diode_type.get() == "n-type" else 32

    def analyze(self):
        """Route analysis based on method selection."""
        try:
            vcol = int(self.volt_entry.get()) - 1
            icol = int(self.curr_entry.get()) - 1
            V = self.df.iloc[:, vcol].values
            I = self.df.iloc[:, icol].values

            # Apply smoothing
            V = self.smooth(V)
            I = self.smooth(I)

            vmin = float(self.vmin_entry.get())
            vmax = float(self.vmax_entry.get())
            T = float(self.temp_entry.get())  # User-defined temperature
            A_star = self.get_A_star()

            if self.method.get() == "Thermionic":
                self.analyze_thermionic(V, I, vmin, vmax, A_star, T)
            else:
                self.analyze_cheung(V, I, vmin, vmax, A_star, T)
        except Exception as e:
            messagebox.showerror("Error", f"Analysis failed:\n{e}")

    def analyze_thermionic(self, V, I, vmin, vmax, A_star, T):
        """Thermionic emission analysis with Rs estimation."""
        # Full-range plot first (I-V semilog)
        self.clear_canvas()
        fig_full, ax_full = plt.subplots()
        ax_full.semilogy(V, np.abs(I), label="I-V (Full Range)")
        ax_full.set_title("I-V Semilog Plot (Full Range)")
        ax_full.set_xlabel("Voltage (V)")
        ax_full.set_ylabel("Current (A, log scale)")
        ax_full.grid(True)
        self.plot_on_canvas(fig_full)

        # Select region for ln(I) vs V
        mask = (V >= vmin) & (V <= vmax)
        Vsel = V[mask]
        Isel = I[mask]

        if len(Vsel) < 2 or np.any(Isel <= 0):
            messagebox.showerror("Error", "Invalid range or negative currents in selected region.")
            return

        lnI = np.log(np.clip(Isel, 1e-12, None))  # Safe log

        # Linear Fit
        slope, intercept = np.polyfit(Vsel, lnI, 1)
        n = q / (slope * k * T)
        I0 = np.exp(intercept)
        phi_b = (k * T / q) * np.log((A_star * Aeff * T ** 2) / I0)

        # R²
        lnI_pred = slope * Vsel + intercept
        ss_res = np.sum((lnI - lnI_pred) ** 2)
        ss_tot = np.sum((lnI - np.mean(lnI)) ** 2)
        r2 = 1 - ss_res / ss_tot

        # Estimate Rs using Cheung-like approach
        dV_dlnI = np.gradient(Vsel) / np.gradient(lnI)
        slope_rs, _ = np.polyfit(Isel, dV_dlnI, 1)
        Rs_thermionic = slope_rs  # Approximate Rs from dV/dlnI vs I

        # Plot: ln(I) vs V
        fig_ln, ax_ln = plt.subplots()
        ax_ln.plot(Vsel, lnI, 'bo-', label='ln(I)')
        ax_ln.plot(Vsel, lnI_pred, 'r--', label='Fit')
        ax_ln.set_title("ln(I) vs V (Selected Range)")
        ax_ln.set_xlabel("Voltage (V)")
        ax_ln.set_ylabel("ln(Current)")
        ax_ln.legend()
        ax_ln.grid(True)
        self.plot_on_canvas(fig_ln)

        # Display Results
        results = [
            ("Ideality Factor (n)", n),
            ("Reverse Saturation Current (I₀, A)", I0),
            ("Barrier Height (Φb, V)", phi_b),
            ("Series Resistance (Rs, Ω)", Rs_thermionic),
            ("Slope", slope),
            ("Intercept", intercept),
            ("R² (Fit Quality)", r2),
            ("Temperature (K)", T)
        ]

        self.display_results(results, list(zip(Vsel, Isel)), method="Thermionic")

    def analyze_cheung(self, V, I, vmin, vmax, A_star, T):
        """Cheung's method with ideality factor, Rs, I₀ and Φb."""
        mask = (V >= vmin) & (V <= vmax)
        Vsel = V[mask]
        Isel = I[mask]

        if len(Vsel) < 5 or np.any(Isel <= 0):
            messagebox.showerror("Error", "Cheung's method requires 5+ points with positive current in forward region.")
            return

        # Step 1: Derivatives
        lnI = np.log(np.clip(Isel, 1e-12, None))
        dV_dlnI = np.gradient(Vsel) / np.gradient(lnI)
        H_I = Vsel - (k * T / q) * lnI

        # Step 2: Fit dV/dlnI = Rs × I + (n·kT/q)
        slope1, intercept1 = np.polyfit(Isel, dV_dlnI, 1)
        Rs = slope1
        n = intercept1 * q / (k * T)

        # Step 3: Fit H(I) = Rs × I + Φb
        slope2, intercept2 = np.polyfit(Isel, H_I, 1)
        phi_b = intercept2  # Direct from intercept of H(I) fit

        # Step 4: I₀ from Cheung expression:
        I0 = np.exp(-q * phi_b / (n * k * T))

        # Rs mismatch warning
        if abs(Rs - slope2) / max(Rs, slope2) > 0.2:
            messagebox.showwarning("Rs Mismatch", "⚠️ Rs1 and Rs2 differ significantly (>20%)")

        # Plot 1: dV/dlnI vs I
        self.clear_canvas()
        fig1, ax1 = plt.subplots()
        ax1.plot(Isel, dV_dlnI, 'bo-', label='dV/dlnI')
        ax1.plot(Isel, slope1 * Isel + intercept1, 'r--', label='Fit')
        ax1.set_title("dV/dlnI vs Current")
        ax1.set_xlabel("Current (A)")
        ax1.set_ylabel("dV/dln(I) (V)")
        ax1.legend()
        ax1.grid(True)
        self.plot_on_canvas(fig1)

        # Plot 2: H(I) vs I
        fig2, ax2 = plt.subplots()
        ax2.plot(Isel, H_I, 'go-', label='H(I)')
        ax2.plot(Isel, slope2 * Isel + intercept2, 'm--', label='Fit')
        ax2.set_title("H(I) vs Current")
        ax2.set_xlabel("Current (A)")
        ax2.set_ylabel("H(I) (V)")
        ax2.legend()
        ax2.grid(True)
        self.plot_on_canvas(fig2)

        # Display Results
        results = [
            ("Ideality Factor (n)", n),
            ("Reverse Saturation Current (I₀, A)", I0),
            ("Barrier Height (Φb, V)", phi_b),
            ("Series Resistance Rs (from dV/dlnI)", Rs),
            ("Series Resistance Rs (from H(I))", slope2),
            ("Slope1", slope1),
            ("Intercept1 (n·kT/q)", intercept1),
            ("Slope2", slope2),
            ("Intercept2 (Φb)", intercept2),
            ("Temperature (K)", T)
        ]

        self.display_results(results, list(zip(Vsel, Isel)), method="Cheung")

    def best_exponential_region(self):
        """Auto-detect best exponential ln(I)-V region with at least 6 points within current Vmin/Vmax range."""
        try:
            vcol = int(self.volt_entry.get()) - 1
            icol = int(self.curr_entry.get()) - 1
            V = self.df.iloc[:, vcol].values
            I = self.df.iloc[:, icol].values

            # Get current Vmin/Vmax from GUI
            current_vmin = float(self.vmin_entry.get())
            current_vmax = float(self.vmax_entry.get())

            # Filter within current range and positive current
            mask = (V >= current_vmin) & (V <= current_vmax) & (I > 0)
            V = V[mask]
            I = I[mask]

            if len(V) < 6:
                messagebox.showerror("Error", "Need at least 6 points in current range")
                return

            lnI = np.log(np.clip(I, 1e-12, None))

            best_r2 = -np.inf
            best_range = (V[0], V[5])  # Initialize with first 6 points

            # Search only within current Vmin/Vmax with minimum 6 points
            for i in range(len(V) - 5):  # Need at least 6 points (i to i+5)
                for j in range(i + 6, len(V) + 1):  # j is exclusive, so i+6 gives 6 points
                    x, y = V[i:j], lnI[i:j]
                    slope, intercept = np.polyfit(x, y, 1)
                    y_pred = slope * x + intercept
                    ss_res = np.sum((y - y_pred) ** 2)
                    ss_tot = np.sum((y - np.mean(y)) ** 2)
                    r2 = 1 - (ss_res / ss_tot) if ss_tot != 0 else 0

                    if r2 > best_r2:
                        best_r2 = r2
                        best_range = (V[i], V[j - 1])

            # Update GUI fields
            self.vmin_entry.delete(0, tk.END)
            self.vmin_entry.insert(0, f"{best_range[0]:.4f}")
            self.vmax_entry.delete(0, tk.END)
            self.vmax_entry.insert(0, f"{best_range[1]:.4f}")

            messagebox.showinfo("Best Region Found",
                                f"Vmin = {best_range[0]:.4f} V\n"
                                f"Vmax = {best_range[1]:.4f} V\n"
                                f"Points = {np.sum((V >= best_range[0]) & (V <= best_range[1]))}\n"
                                f"R² = {best_r2:.5f}")
        except Exception as e:
            messagebox.showerror("Best Region Error", str(e))


class SurfaceStateAnalyzer(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.controller = controller

        # Constants
        self.q = 1.602e-19  # Electron charge (C)
        self.k = 1.381e-23  # Boltzmann constant (J/K)

        # Data storage
        self.df = None
        self.results_df = None

        self.setup_gui()

    def setup_gui(self):
        """Setup the main GUI components"""
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # File Input Section
        file_frame = ttk.LabelFrame(main_frame, text="File Input", padding="10")
        file_frame.pack(fill=tk.X, pady=5)

        ttk.Button(file_frame, text="Load Data File", command=self.load_file).pack(side=tk.LEFT)
        self.file_label = ttk.Label(file_frame, text="No file loaded")
        self.file_label.pack(side=tk.LEFT, padx=10)

        # Column Selection
        col_frame = ttk.LabelFrame(main_frame, text="Data Columns", padding="10")
        col_frame.pack(fill=tk.X, pady=5)

        ttk.Label(col_frame, text="Voltage Column:").grid(row=0, column=0, sticky="e")
        self.v_col = ttk.Entry(col_frame, width=5)
        self.v_col.grid(row=0, column=1, sticky="w")

        ttk.Label(col_frame, text="Current Column:").grid(row=1, column=0, sticky="e")
        self.i_col = ttk.Entry(col_frame, width=5)
        self.i_col.grid(row=1, column=1, sticky="w")

        # Parameters Input
        param_frame = ttk.LabelFrame(main_frame, text="Parameters", padding="10")
        param_frame.pack(fill=tk.X, pady=5)

        params = [
            ("Reverse Saturation Current (I₀, A):", "i0"),
            ("Barrier Height (Φb, eV):", "phi_b"),
            ("Insulator Permittivity (εᵢ):", "eps_i"),
            ("Semiconductor Permittivity (εₛ):", "eps_s"),
            ("Depletion Width (W_D, cm):", "w_d"),
            ("Insulator Thickness (δ, cm):", "delta"),
            ("Temperature (T, K):", "temp")
        ]

        for i, (label, name) in enumerate(params):
            ttk.Label(param_frame, text=label).grid(row=i, column=0, sticky="e")
            entry = ttk.Entry(param_frame)
            entry.grid(row=i, column=1, sticky="w")
            setattr(self, f"{name}_entry", entry)

        # Set some default values
        self.i0_entry.insert(0, "1e-12")
        self.phi_b_entry.insert(0, "0.8")
        self.eps_i_entry.insert(0, "12")
        self.eps_s_entry.insert(0, "11.7")
        self.w_d_entry.insert(0, "1e-5")
        self.delta_entry.insert(0, "1e-7")
        self.temp_entry.insert(0, "300")

        # Action Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=10)

        ttk.Button(btn_frame, text="Calculate Nss", command=self.calculate_nss).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Export Results", command=self.export_results).pack(side=tk.LEFT, padx=5)

        # Results Display
        results_frame = ttk.LabelFrame(main_frame, text="Results", padding="10")
        results_frame.pack(fill=tk.BOTH, expand=True)

        # Treeview for results
        self.tree = ttk.Treeview(results_frame, columns=("V", "I", "n", "Phi_eff", "Nss"), show="headings")

        # Define headings
        self.tree.heading("V", text="Voltage (V)")
        self.tree.heading("I", text="Current (A)")
        self.tree.heading("n", text="Ideality Factor n(V)")
        self.tree.heading("Phi_eff", text="Effective Φb (eV)")
        self.tree.heading("Nss", text="Nss (eV⁻¹cm⁻²)")

        # Set column widths
        for col in ("V", "I", "n", "Phi_eff", "Nss"):
            self.tree.column(col, width=120, anchor="center")

        # Add scrollbars
        yscroll = ttk.Scrollbar(results_frame, orient="vertical", command=self.tree.yview)
        xscroll = ttk.Scrollbar(results_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)

        # Layout
        self.tree.grid(row=0, column=0, sticky="nsew")
        yscroll.grid(row=0, column=1, sticky="ns")
        xscroll.grid(row=1, column=0, sticky="ew")

        # Configure grid weights
        results_frame.grid_rowconfigure(0, weight=1)
        results_frame.grid_columnconfigure(0, weight=1)

    def load_file(self):
        """Load data from Excel or CSV file"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Excel files", "*.xlsx *.xls"), ("CSV files", "*.csv")]
        )

        if file_path:
            try:
                if file_path.endswith('.csv'):
                    self.df = pd.read_csv(file_path)
                else:
                    self.df = pd.read_excel(file_path)

                # Convert to numeric and drop non-numeric rows
                self.df = self.df.apply(pd.to_numeric, errors='coerce').dropna()
                self.file_label.config(text=file_path.split('/')[-1])
                messagebox.showinfo("Success", f"Loaded {len(self.df)} data points")

            except Exception as e:
                messagebox.showerror("Error", f"Failed to load file:\n{str(e)}")

    def calculate_nss(self):
        """Calculate surface state density for all data points"""
        if self.df is None:
            messagebox.showerror("Error", "Please load a data file first")
            return

        try:
            # Get column indices
            v_col = int(self.v_col.get()) - 1
            i_col = int(self.i_col.get()) - 1

            # Get parameters with proper unit conversions
            params = {
                'i0': float(self.i0_entry.get()),
                'phi_b': float(self.phi_b_entry.get()),
                'eps_i': float(self.eps_i_entry.get()) * 8.854e-14,  # Convert to F/cm
                'eps_s': float(self.eps_s_entry.get()) * 8.854e-14,  # Convert to F/cm
                'w_d': float(self.w_d_entry.get()),
                'delta': float(self.delta_entry.get()),
                'temp': float(self.temp_entry.get())
            }

            # Extract data
            V = pd.to_numeric(self.df.iloc[:, v_col], errors='coerce').dropna().values
            I = pd.to_numeric(self.df.iloc[:, i_col], errors='coerce').dropna().values

            if len(V) == 0 or len(I) == 0:
                raise ValueError("No valid numeric data in selected columns")

            # Calculate results
            results = []
            for v, i in zip(V, I):
                if i <= 0:
                    continue  # Skip negative/zero currents

                # Calculate ideality factor (convert k to eV/K)
                k_ev = 8.617e-5  # Boltzmann constant in eV/K
                n = (v) / (k_ev * params['temp'] * np.log(i / params['i0']))

                # Calculate surface state density
                term1 = (params['eps_i'] / params['delta']) * (n - 1)
                term2 = params['eps_s'] / params['w_d']
                Nss = (1 / self.q) * (term1 - term2)

                # Convert to more reasonable units (cm⁻²eV⁻¹)
                Nss = Nss * 1e-4  # Adjust scale if needed

                results.append((v, i, n, params['phi_b'], abs(Nss)))  # Use abs() instead of max()

            # Store and display results
            self.results_df = pd.DataFrame(results,
                                           columns=["Voltage (V)", "Current (A)",
                                                    "Ideality Factor", "Effective Φb (eV)",
                                                    "Nss (eV⁻¹cm⁻²)"])

            self.display_results()
            messagebox.showinfo("Success", f"Calculated Nss for {len(results)} data points")

        except ValueError as ve:
            messagebox.showerror("Input Error", f"Invalid input:\n{str(ve)}")
        except Exception as e:
            messagebox.showerror("Calculation Error", f"Failed to calculate Nss:\n{str(e)}")

    def display_results(self):
        """Display results in the treeview"""
        # Clear previous results
        for item in self.tree.get_children():
            self.tree.delete(item)

        # Add new results
        for _, row in self.results_df.iterrows():
            self.tree.insert("", "end", values=tuple(row))

    def export_results(self):
        """Export results to Excel file"""
        if self.results_df is None:
            messagebox.showerror("Error", "No results to export")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("CSV files", "*.csv")]
        )

        if file_path:
            try:
                if file_path.endswith('.csv'):
                    self.results_df.to_csv(file_path, index=False)
                else:
                    self.results_df.to_excel(file_path, index=False)

                messagebox.showinfo("Success", f"Results exported to:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Export Error", f"Failed to export results:\n{str(e)}")


# ----------------------------- Launch App -----------------------------
if __name__ == "__main__":
    root = tk.Tk()
    app = MainApplication(root)
    root.mainloop()